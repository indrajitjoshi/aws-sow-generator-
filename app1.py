import streamlit as st
from datetime import date
import io, re, os, requests
import pandas as pd

# ================= AWS DIAGRAM IMPORTS =================
from diagrams import Diagram, Cluster
from diagrams.aws.general import User
from diagrams.aws.compute import Lambda, EC2
from diagrams.aws.ml import Bedrock, Textract, Rekognition
from diagrams.aws.storage import S3
from diagrams.aws.database import OpenSearch, RDS
from diagrams.aws.analytics import Kinesis
from diagrams.aws.iot import IotCore
# ======================================================

# ================= PAGE CONFIG =================
st.set_page_config(
    page_title="GenAI SOW Architect",
    layout="wide",
    page_icon="📄",
    initial_sidebar_state="expanded"
)

# ================= ARCHITECTURE PATTERN MAP =================
ARCH_PATTERN_MAP = {
    "Multi Agent Store Advisor": "AGENTIC_RAG",
    "Agentic AI L1 Support": "AGENTIC_RAG",
    "Sales Co-Pilot": "AGENTIC_RAG",
    "Research Co-Pilot": "AGENTIC_RAG",
    "SOP Creation": "AGENTIC_RAG",
    "Multi-agent e-KYC & Onboarding": "AGENTIC_RAG",

    "Intelligent Search": "RAG_TEXT",
    "Document / Report Audit": "RAG_TEXT",
    "RBI Circular Scraping & Insights Bot": "RAG_TEXT",
    "Customer Review Analysis": "RAG_TEXT",
    "Cost, Margin Visibility & Insights using LLM": "RAG_TEXT",

    "Virtual Data Analyst (Text to SQL)": "TEXT_TO_SQL",

    "Recommendation": "RECOMMENDER",
    "AI Agents Demand Forecasting": "RECOMMENDER",
    "AI Agents Based Pricing Module": "RECOMMENDER",
    "AI Trend Simulator": "RECOMMENDER",

    "Banner Audit using LLM": "VISION_LLM",
    "Image Enhancement": "VISION_LLM",
    "Virtual Try-On": "VISION_LLM",
    "Visual Inspection": "VISION_LLM",

    "Multilingual Call Analysis": "VOICE_AI",
    "Multilingual Voice Bot": "VOICE_AI",

    "AIoT based CCTV Surveillance": "IOT_STREAM",

    "Product Listing Standardization": "CONTENT_GEN",
    "Product Copy Generator": "CONTENT_GEN"
}

# ================= ARCHITECTURE SPEC =================
def generate_architecture_spec(solution):
    pattern = ARCH_PATTERN_MAP.get(solution, "RAG_TEXT")

    spec = {
        "title": f"{solution} – AWS Reference Architecture",
        "pattern": pattern,
        "security_boundary": "Customer VPC",
        "vector_store": pattern in ["RAG_TEXT", "AGENTIC_RAG"],
        "data_sources": ["S3", "Textract"],
        "extras": []
    }

    if pattern == "VISION_LLM":
        spec["extras"].append("Rekognition")
    if pattern == "VOICE_AI":
        spec["extras"].append("Kinesis")
    if pattern == "IOT_STREAM":
        spec["extras"] += ["IoT Core", "Kinesis"]

    return spec

# ================= DIAGRAM RENDERER =================
def render_architecture_diagram(spec):
    fname = f"architecture_{spec['title'].replace(' ', '_')}"
    with Diagram(spec["title"], filename=fname, show=False, direction="LR"):
        user = User("User")
        with Cluster("AWS Cloud"):
            with Cluster(spec["security_boundary"]):
                ui = EC2("Streamlit UI")
                orchestrator = Lambda("Query + Agent Orchestration")
                llm = Bedrock("Amazon Bedrock\n(Mistral)")
                components = []

                if spec["vector_store"]:
                    components.append(OpenSearch("Vector DB"))
                components += [S3("Documents"), Textract("OCR")]

                if "Rekognition" in spec["extras"]:
                    components.append(Rekognition("Vision AI"))
                if "IoT Core" in spec["extras"]:
                    components.append(IotCore("IoT Ingest"))
                if "Kinesis" in spec["extras"]:
                    components.append(Kinesis("Streaming"))

                db = RDS("Application DB")

        user >> ui >> orchestrator >> llm
        for c in components:
            orchestrator >> c >> orchestrator
        orchestrator >> db
        orchestrator >> ui

    return f"{fname}.png"

# ================= DOCX GENERATOR =================
def create_docx(text, diagram_path, solution_name, doc_date):
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading(solution_name, 0)
    doc.add_paragraph("Scope of Work Document")
    doc.add_paragraph(doc_date.strftime("%d %B %Y"))
    doc.add_page_break()

    for line in text.split("\n"):
        doc.add_paragraph(line)

    doc.add_page_break()
    doc.add_heading("4 SOLUTION ARCHITECTURE / ARCHITECTURAL DIAGRAM", 1)
    if diagram_path and os.path.exists(diagram_path):
        doc.add_picture(diagram_path, width=Inches(6.5))

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ================= SESSION =================
if "sow" not in st.session_state:
    st.session_state.sow = ""

# ================= SIDEBAR =================
with st.sidebar:
    st.title("SOW Architect")
    api_key = st.text_input("Gemini API Key", type="password")

    solution = st.selectbox("Solution Type", list(ARCH_PATTERN_MAP.keys()))
    engagement = st.selectbox("Engagement Type", [
        "Proof of Concept (PoC)", "Pilot", "MVP",
        "Production Rollout", "Assessment / Discovery"
    ])
    industry = st.selectbox("Industry", [
        "Retail / E-commerce", "BFSI", "Healthcare", "Manufacturing",
        "Telecom", "Government", "Other"
    ])
    timeline = st.text_input("Timeline", "4 Weeks")

# ================= MAIN UI =================
st.title("🚀 GenAI Scope of Work Architect")

objective = st.text_area(
    "2.1 Define the core business objective:",
    height=120
)

if st.button("✨ Generate SOW + Architecture", use_container_width=True):
    if not api_key or not objective:
        st.error("API Key and Objective are required")
    else:
        with st.spinner("Generating SOW and Architecture..."):
            # ---- Gemini Call ----
            prompt = f"""
Generate a COMPLETE enterprise Scope of Work (SOW).

MANDATORY STRUCTURE:
1 TABLE OF CONTENTS
2 PROJECT OVERVIEW
2.1 OBJECTIVE
2.2 PROJECT SPONSOR(S) / STAKEHOLDER(S) / PROJECT TEAM
2.3 ASSUMPTIONS & DEPENDENCIES
2.4 PROJECT SUCCESS CRITERIA
3 SCOPE OF WORK – TECHNICAL PROJECT PLAN
4 SOLUTION ARCHITECTURE / ARCHITECTURAL DIAGRAM
5 RESOURCES & COST ESTIMATES

INPUTS:
Solution: {solution}
Industry: {industry}
Engagement Type: {engagement}
Timeline: {timeline}
Objective: {objective}

Rules:
- No markdown symbols
- Plain text only
- Strict numbering
"""

            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={api_key}"
            payload = {
                "contents": [{"parts": [{"text": prompt}]}]
            }

            res = requests.post(url, json=payload)
            st.session_state.sow = res.json()["candidates"][0]["content"]["parts"][0]["text"]

            # ---- Architecture ----
            spec = generate_architecture_spec(solution)
            diagram_path = render_architecture_diagram(spec)
            st.session_state.diagram = diagram_path

# ================= OUTPUT =================
if st.session_state.sow:
    st.divider()
    st.header("📄 Generated SOW")
    st.text_area("SOW Content", st.session_state.sow, height=500)

if "diagram" in st.session_state:
    st.divider()
    st.header("🧱 Solution Architecture")
    st.image(st.session_state.diagram, use_column_width=True)

    docx = create_docx(
        st.session_state.sow,
        st.session_state.diagram,
        solution,
        date.today()
    )

    st.download_button(
        "📥 Download SOW (.docx)",
        data=docx,
        file_name=f"SOW_{solution.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
