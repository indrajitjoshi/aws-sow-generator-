import streamlit as st
from datetime import date
import io
import re
import os
import time 
import requests
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- FILE PATHING & DIAGRAM MAPPING ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ASSETS_DIR = os.path.join(BASE_DIR, "diagrams")
if not os.path.exists(ASSETS_DIR):
    os.makedirs(ASSETS_DIR, exist_ok=True)

# Static assets (ensure these files exist in the /diagrams folder)
AWS_PN_LOGO = os.path.join(ASSETS_DIR, "aws partner logo.jpg")
ONETURE_LOGO = os.path.join(ASSETS_DIR, "oneture logo1.jpg")
AWS_ADV_LOGO = os.path.join(ASSETS_DIR, "aws advanced logo1.jpg")

# Mapped Infra Costs
SOW_COST_TABLE_MAP = { 
    "L1 Support Bot POC SOW": { "poc_cost": "3,536.40 USD" }, 
    "Beauty Advisor POC SOW": { 
        "poc_cost": "4,525.66 USD + 200 USD (Amazon Bedrock Cost) = 4,725.66", 
        "prod_cost": "4,525.66 USD + 1,175.82 USD (Amazon Bedrock Cost) = 5,701.48" 
    }, 
    "Ready Search POC Scope of Work Document":{ "poc_cost": "2,641.40 USD" }, 
    "AI based Image Enhancement POC SOW": { "poc_cost": "2,814.34 USD" }, 
    "AI based Image Inspection POC SOW": { "poc_cost": "3,536.40 USD" }, 
    "Gen AI for SOP POC SOW": { "poc_cost": "2,110.30 USD" }, 
    "Project Scope Document": { "prod_cost": "2,993.60 USD" }, 
    "Gen AI Speech To Speech": { "prod_cost": "2,124.23 USD" }, 
    "PoC Scope Document": { "amazon_bedrock": "1,000 USD", "total": "$ 3,150" }
}

# Mapped Calculator Links
CALCULATOR_LINKS = {
    "L1 Support Bot POC SOW": "https://calculator.aws/#/estimate?id=211ea64cba5a8f5dc09805f4ad1a1e598ef5238b",
    "Ready Search POC Scope of Work Document": "https://calculator.aws/#/estimate?id=f8bc48f1ae566b8ea1241994328978e7e86d3490",
    "AI based Image Enhancement POC SOW": "https://calculator.aws/#/estimate?id=9a3e593b92b796acecf31a78aec17d7eb957d1e5",
    "Beauty Advisor POC SOW": "https://calculator.aws/#/estimate?id=3f89756a35f7bac7b2cd88d95f3e9aba9be9b0eb",
    "Beauty Advisor Production": "https://calculator.aws/#/estimate?id=4d7f092e819c799f680fd14f8de3f181f565c48e",
    "AI based Image Inspection POC SOW": "https://calculator.aws/#/estimate?id=72c56f93b0c0e101d67a46af4f4fe9886eb93342",
    "Gen AI for SOP POC SOW": "https://calculator.aws/#/estimate?id=c21e9b242964724bf83556cfeee821473bb935d1",
    "Project Scope Document": "https://calculator.aws/#/estimate?id=37339d6e34c73596559fe09ca16a0ac2ec4c4252",
    "Gen AI Speech To Speech": "https://calculator.aws/#/estimate?id=8444ae26e6d61e5a43e8e743578caa17fd7f3e69",
    "PoC Scope Document": "https://calculator.aws/#/estimate?id=420ed9df095e7824a144cb6c0e9db9e7ec3c4153"
}

# Image Mapping (using filenames)
SOW_DIAGRAM_MAP = {
    "L1 Support Bot POC SOW": "L1 Support Bot POC SOW.png",
    "Beauty Advisor POC SOW": "Beauty Advisor POC SOW.png",
    "Ready Search POC Scope of Work Document": "Ready Search POC Scope of Work Document.png",
    "AI based Image Enhancement POC SOW": "AI based Image Enhancement POC SOW.png",
    "AI based Image Inspection POC SOW": "AI based Image Inspection POC SOW.png",
    "Gen AI for SOP POC SOW": "Gen AI for SOP POC SOW.png",
    "Project Scope Document": "Project Scope Document.png",
    "Gen AI Speech To Speech": "Gen AI Speech To Speech.png",
    "PoC Scope Document": "PoC Scope Document.png"
}

# --- CONFIGURATION ---
st.set_page_config(
    page_title="GenAI SOW Architect", 
    layout="wide", 
    page_icon="📄",
    initial_sidebar_state="expanded"
)

# Custom CSS for an Enterprise UI
st.markdown("""
    <style>
    .main { background-color: #f8fafc; }
    .stButton>button { border-radius: 8px; font-weight: 600; }
    .stTextArea textarea { border-radius: 10px; }
    .stTextInput input { border-radius: 8px; }
    .block-container { padding-top: 1.5rem; }
    .sow-preview {
        background-color: white;
        padding: 40px;
        border-radius: 12px;
        border: 1px solid #e2e8f0;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        line-height: 1.7;
        color: #1e293b;
        box-shadow: 0 4px 6px -1px rgb(0 0 0 / 0.1);
    }
    h1, h2, h3 { color: #0f172a; }
    .stTabs [data-baseweb="tab-list"] { gap: 24px; }
    .stTabs [data-baseweb="tab"] { height: 50px; white-space: pre-wrap; font-weight: 600; }
    .stakeholder-header { 
        background-color: #f1f5f9; 
        padding: 8px 12px; 
        border-radius: 6px; 
        margin-bottom: 10px; 
        font-weight: bold;
        color: #334155;
        border-left: 4px solid #3b82f6;
    }
    </style>
    """, unsafe_allow_html=True)

# --- UTILITIES ---

def safe_add_picture(doc, image_path, width):
    """Safely adds a picture to the docx document."""
    try:
        if not image_path or not os.path.exists(image_path):
            return False
        with Image.open(image_path) as img:
            img.verify()
        doc.add_picture(image_path, width=width)
        return True
    except Exception as e:
        print(f"[DOCX IMAGE SKIPPED] {image_path} -> {e}")
        return False

def add_hyperlink(paragraph, text, url):
    import docx.oxml.shared
    import docx.opc.constants
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '0000FF')
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(c)
    rPr.append(u)
    new_run.append(rPr)
    t = docx.oxml.shared.OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def add_poc_calculation_table(doc):
    doc.add_paragraph("The above numbers are calculated basis the following:")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Particulars", "Value (in Dollar)", "Remarks"
    data = [
        ("Number of documents", "200", "Assuming 5 interactions"),
        ("Input Tokens per document", "10,00,000", ""),
        ("Input Token Cost / 1k", "0", "Anthropic Claude 3 Sonnet"),
        ("Total Input Cost", "600", ""),
        ("Output Tokens", "50,000", ""),
        ("Total Output Cost", "150", ""),
        ("Total Cost", "750", ""),
        ("Tokens for Embedding", "2,50,00,00,000", ""),
        ("Total Embedding Cost", "250", ""),
        ("Total Cost per month", "1,000", "")
    ]
    for row in data:
        cells = table.add_row().cells
        for i, val in enumerate(row): cells[i].text = val

def add_infra_cost_table(doc, sow_type_name, text_content):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    cost_data = SOW_COST_TABLE_MAP.get(sow_type_name)
    if not cost_data: return

    calc_url = CALCULATOR_LINKS.get(sow_type_name, "https://calculator.aws/#/")
    if sow_type_name == "Beauty Advisor POC SOW" and "Production Development" in text_content:
        calc_url = CALCULATOR_LINKS["Beauty Advisor Production"]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "System", "Infra Cost / month", "AWS Calculator Cost"

    rows_to_add = []
    if "poc_cost" in cost_data: rows_to_add.append(("POC", cost_data["poc_cost"]))
    if "prod_cost" in cost_data: rows_to_add.append(("Production", cost_data["prod_cost"]))
    if "amazon_bedrock" in cost_data: rows_to_add.append(("Amazon Bedrock", cost_data["amazon_bedrock"]))
    if "total" in cost_data: rows_to_add.append(("Total", cost_data["total"]))

    for label, cost in rows_to_add:
        r = table.add_row().cells
        r[0].text, r[1].text = label, cost
        p = r[2].paragraphs[0]
        add_hyperlink(p, "Estimate", calc_url)

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if sow_type_name == "PoC Scope Document":
        doc.add_paragraph("")
        add_poc_calculation_table(doc)

def create_docx_logic(text_content, branding_info, sow_type_name):
    doc = Document()
    header_patterns = {
        "1": "1 TABLE OF CONTENTS", "2": "2 PROJECT OVERVIEW", "3": "3 ASSUMPTIONS", 
        "4": "4 PROJECT SUCCESS", "5": "5 SCOPE OF WORK", "6": "6 SOLUTION ARCHITECTURE",
        "7": "7 PERFORMANCE", "8": "8 COST ESTIMATION", "9": "9 RESOURCES", "10": "10 FINAL"
    }
    rendered_sections = {k: False for k in header_patterns.keys()}

    # COVER PAGE
    p_top = doc.add_paragraph()
    safe_add_picture(doc, AWS_PN_LOGO, Inches(1.6))
    doc.add_paragraph("\n" * 3)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(branding_info['sow_name'])
    run.font.size, run.bold = Pt(26), True
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.add_run("Scope of Work Document").font.size = Pt(14)
    doc.add_paragraph("\n" * 4)

    logo_table = doc.add_table(rows=1, cols=3)
    logo_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = logo_table.rows[0].cells[0]
    if branding_info.get("customer_logo_bytes"):
        cell.paragraphs[0].add_run().add_picture(io.BytesIO(branding_info["customer_logo_bytes"]), width=Inches(1.8))
    else: cell.paragraphs[0].add_run("Customer Logo").bold = True
    safe_add_picture(doc, ONETURE_LOGO, Inches(2.2)) # Middle
    safe_add_picture(doc, AWS_ADV_LOGO, Inches(1.8)) # Right

    doc.add_page_break()

    # CONTENT
    lines = text_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line: i += 1; continue
        line_clean = re.sub(r'\*+', '', line).strip()
        clean_text = re.sub(r'^#+\s*', '', line_clean).strip()
        upper_text = clean_text.upper()

        current_header_id = next((h_id for h_id, pattern in header_patterns.items() if pattern in upper_text), None)

        if current_header_id:
            if not rendered_sections.get(current_header_id):
                if current_header_id == "2": doc.add_page_break()
                doc.add_heading(clean_text, level=1)
                rendered_sections[current_header_id] = True
                
                if current_header_id == "6":
                    filename = SOW_DIAGRAM_MAP.get(sow_type_name)
                    if filename:
                        path = os.path.join(ASSETS_DIR, filename)
                        if safe_add_picture(doc, path, Inches(5.8)):
                            cap = doc.add_paragraph(f"{sow_type_name} – Architecture Diagram")
                            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if current_header_id == "8":
                    add_infra_cost_table(doc, sow_type_name, text_content)
            i += 1; continue

        if line.startswith('## '): doc.add_heading(clean_text, level=2)
        elif line.startswith('### '): doc.add_heading(clean_text, level=3)
        elif line.startswith(('- ', '* ')): doc.add_paragraph(line[2:], style="List Bullet")
        else: doc.add_paragraph(line)
        i += 1
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def call_gemini_with_retry(api_key, payload):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={api_key}"
    retries = 5
    for attempt in range(retries):
        try:
            res = requests.post(url, json=payload)
            if res.status_code == 200: return res, None
            if res.status_code in [503, 429]:
                time.sleep(2**attempt); continue
            return None, f"API Error {res.status_code}: {res.text}"
        except: time.sleep(2**attempt)
    return None, "Model overloaded after retries."

# --- INITIALIZATION ---
init_keys = [
    "engagement_type", "success_dimensions", "customer_dependencies", "data_types",
    "data_characteristics", "key_assumptions", "other_assumptions", "user_validation_required",
    "compute_orchestration", "genai_services", "storage_services", "ui_layer",
    "performance_expectation", "security_compliance", "cost_ownership", "deliverables",
    "post_poc_next_steps", "poc_duration", "phase_breakdown"
]
default_vals = {
    "engagement_type": "Proof of Concept (PoC)", "success_dimensions": [], "customer_dependencies": [],
    "data_types": [], "data_characteristics": {}, "key_assumptions": [], "other_assumptions": "",
    "user_validation_required": "Yes – customer validation required", "compute_orchestration": "AWS Lambda + Step Functions",
    "genai_services": ["Amazon Bedrock (LLM inference)"], "storage_services": ["Amazon S3"],
    "ui_layer": "Streamlit on S3", "performance_expectation": "Batch", "security_compliance": [],
    "cost_ownership": "Funded by Customer", "deliverables": [], "post_poc_next_steps": [],
    "poc_duration": "4 weeks", "phase_breakdown": {"Infra setup": "", "Core workflows": "", "Testing & validation": "", "Demo & feedback": ""}
}
for k in init_keys:
    if k not in st.session_state: st.session_state[k] = default_vals[k]

if 'generated_sow' not in st.session_state: st.session_state.generated_sow = ""
if 'stakeholders' not in st.session_state:
    import pandas as pd
    st.session_state.stakeholders = {
        "Partner": pd.DataFrame([{"Name": "Gaurav Kankaria", "Title": "Head of Analytics & ML", "Email": "gaurav.kankaria@oneture.com"}]),
        "Customer": pd.DataFrame([{"Name": "Cheten Dev", "Title": "Head of Product Design", "Email": "cheten.dev@nykaa.com"}]),
        "AWS": pd.DataFrame([{"Name": "Anubhav Sood", "Title": "AWS Account Executive", "Email": "anbhsood@amazon.com"}]),
        "Escalation": pd.DataFrame([{"Name": "Omkar Dhavalikar", "Title": "AI/ML Lead", "Email": "omkar.dhavalikar@oneture.com"}])
    }

def clear_sow(): st.session_state.generated_sow = ""

# --- SIDEBAR ---
with st.sidebar:
    st.title("SOW Architect")
    api_key = st.text_input("Gemini API Key", type="password")
    st.divider()
    st.header("1. Project Intake")
    st.radio("Engagement Type", ["Proof of Concept (PoC)", "Pilot", "MVP", "Production Rollout", "Assessment / Discovery"], key="engagement_type")
    sow_type_options = list(SOW_COST_TABLE_MAP.keys())
    selected_sow_name = st.selectbox("Scope of Work Type", sow_type_options)
    industry = st.selectbox("Industry", ["Retail / E-commerce", "BFSI", "Manufacturing", "Healthcare", "Other"])
    duration = st.text_input("Timeline", "4 Weeks")
    if st.button("Reset All Fields"): st.session_state.generated_sow = ""; st.rerun()

# --- MAIN UI ---
st.title("GenAI Scope of Work Architect")
st.header("Cover Page Branding")
customer_logo = st.file_uploader("Upload Customer Logo", type=["png", "jpg", "jpeg"])
doc_date = st.date_input("Document Date", date.today())
st.divider()

st.header("2. Objectives & Stakeholders")
objective = st.text_area("Define business objective:", placeholder="e.g., WIMO Bot development...", height=100)
outcomes = st.multiselect("Key Outcomes:", ["Reduce manual effort", "Improve accuracy", "Cost reduction", "Better CX"])
st.divider()

col_t1, col_t2 = st.columns(2)
with col_t1:
    st.markdown('<div class="stakeholder-header">Partner Sponsor</div>', unsafe_allow_html=True)
    st.session_state.stakeholders["Partner"] = st.data_editor(st.session_state.stakeholders["Partner"], num_rows="dynamic", key="ed_p")
    st.markdown('<div class="stakeholder-header">AWS Sponsor</div>', unsafe_allow_html=True)
    st.session_state.stakeholders["AWS"] = st.data_editor(st.session_state.stakeholders["AWS"], num_rows="dynamic", key="ed_a")
with col_t2:
    st.markdown('<div class="stakeholder-header">Customer Sponsor</div>', unsafe_allow_html=True)
    st.session_state.stakeholders["Customer"] = st.data_editor(st.session_state.stakeholders["Customer"], num_rows="dynamic", key="ed_c")
    st.markdown('<div class="stakeholder-header">Escalation Contacts</div>', unsafe_allow_html=True)
    st.session_state.stakeholders["Escalation"] = st.data_editor(st.session_state.stakeholders["Escalation"], num_rows="dynamic", key="ed_e")

st.divider()
st.header("3. Intake Details")
st.multiselect("Customer Dependencies:", ["Sample data", "API access", "VPC access", "SME availability"], key="customer_dependencies")
st.multiselect("Data Types:", ["Images", "Text", "PDFs", "Audio"], key="data_types")
st.multiselect("Key Assumptions:", ["PoC only", "Limited volume", "Manual review acceptable"], key="key_assumptions")
st.text_area("Other assumptions:", key="other_assumptions")

st.header("4. Success & Logic")
st.multiselect("Success Dimensions:", ["Accuracy", "Latency", "Usability", "Cost efficiency"], key="success_dimensions")
st.radio("Validation Required?", ["Yes – customer validation required", "No – internal validation sufficient"], key="user_validation_required")

st.header("6. Technical Choices")
st.radio("Orchestration:", ["AWS Lambda", "Step Functions", "Hybrid"], key="compute_orchestration")
st.multiselect("GenAI Services:", ["Amazon Bedrock", "SageMaker", "Rekognition", "Textract"], key="genai_services")
st.multiselect("Storage:", ["Amazon S3", "DynamoDB", "OpenSearch", "RDS"], key="storage_services")
st.radio("UI Layer:", ["Streamlit on S3", "CloudFront + Static UI", "API only"], key="ui_layer")

st.header("7. Ops & Timeline")
st.selectbox("Performance:", ["Batch", "Near real-time", "Real-time"], key="performance_expectation")
st.multiselect("Security:", ["IAM-based access", "Encryption", "VPC deployment"], key="security_compliance")
st.radio("Cost Ownership:", ["Funded by AWS", "Funded by Partner", "Funded by Customer"], key="cost_ownership")

# GENERATION
if st.button("✨ Generate SOW Document", type="primary", use_container_width=True):
    if not api_key: st.warning("Enter API Key")
    elif not objective: st.error("Objective required")
    else:
        with st.spinner("Generating..."):
            def get_md(df): return df.to_markdown(index=False)
            prompt = f"""
            Generate a COMPLETE formal enterprise SOW for {selected_sow_name} in {industry}.
            STRUCTURE: 1 TOC, 2 OVERVIEW, 3 ASSUMPTIONS, 4 SUCCESS, 5 SCOPE, 6 ARCHITECTURE, 7 PERFORMANCE, 8 COST, 9 RESOURCES, 10 FINAL.
            CONTEXT: Objective: {objective}. Type: {st.session_state.engagement_type}.
            STAKEHOLDERS:
            {get_md(st.session_state.stakeholders['Partner'])}
            {get_md(st.session_state.stakeholders['Customer'])}
            """
            payload = {"contents": [{"parts": [{"text": prompt}]}], "systemInstruction": {"parts": [{"text": "Solutions Architect. Formal tone."}]}}
            res, err = call_gemini_with_retry(api_key, payload)
            if res:
                st.session_state.generated_sow = res.json()['candidates'][0]['content']['parts'][0]['text']
                st.balloons()
            else: st.error(err)

# REVIEW & EXPORT
if st.session_state.generated_sow:
    st.divider()
    t_edit, t_preview = st.tabs(["✍️ Editor", "📄 Preview"])
    with t_edit:
        st.session_state.generated_sow = st.text_area("Modify:", st.session_state.generated_sow, height=500, key="editor")
    with t_preview:
        st.markdown(f'<div class="sow-preview">', unsafe_allow_html=True)
        st.markdown(st.session_state.generated_sow)
        st.divider()
        st.subheader("Architecture Preview")
        fname = SOW_DIAGRAM_MAP.get(selected_sow_name)
        if fname:
            fpath = os.path.join(ASSETS_DIR, fname)
            if os.path.exists(fpath): st.image(fpath, caption=f"{selected_sow_name} Diagram")
            else: st.warning(f"Diagram missing: {fname}")
        st.markdown('</div>', unsafe_allow_html=True)

    if st.button("💾 Prepare Microsoft Word Document"):
        branding = {"sow_name": selected_sow_name, "customer_logo_bytes": customer_logo.getvalue() if customer_logo else None, "doc_date_str": doc_date.strftime("%d %B %Y")}
        docx_data = create_docx_logic(st.session_state.generated_sow, branding, selected_sow_name)
        st.download_button("📥 Download .docx", docx_data, file_name=f"SOW_{selected_sow_name.replace(' ','_')}.docx")
